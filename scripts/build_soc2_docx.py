import os
import sys
import subprocess

# Self-install python-docx if not present
try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
except ImportError:
    print("python-docx is not installed. Installing it now...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    """Set background color of a table cell."""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner margins (padding) of a table cell in dxa (1/20 of a pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_styled(doc, text, level, space_before=12, space_after=6):
    """Add a heading with custom spacing and colors."""
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after = Pt(space_after)
    h.paragraph_format.keep_with_next = True
    
    # Customize fonts
    run = h.runs[0]
    run.font.name = 'Segoe UI'
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(29, 53, 87) # Slate blue (#1d3557)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(163, 58, 33) # Rust red (#a33a21)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.italic = True
        run.font.color.rgb = RGBColor(80, 80, 80)
    return h

def main():
    print("Starting generation of SOC 2 Internal Audit DOCX...")
    
    doc = docx.Document()
    
    # Page Margins (1 inch all around)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(24, 21, 17) # Soft black (#181511)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)

    # 1. Document Title Page / Header Block
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("Demystifying SOC 2: The Definitive Internal Audit Guide & Readiness Checklist")
    run_title.font.name = 'Segoe UI'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(29, 53, 87)

    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(18)
    run_sub = subtitle_p.add_run("Moving from Compliance Box-Checking to Strategic Risk Governance")
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(163, 58, 33)

    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_after = Pt(24)
    run_meta = meta_p.add_run("Prepared by: Majid Mumtaz, CIA, ACA, FCCA\nDirector of Internal Audit & Governance Advisor")
    run_meta.font.size = Pt(10)
    run_meta.font.bold = True
    run_meta.font.color.rgb = RGBColor(100, 100, 100)

    # Divider line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(18)
    r_div = p_div.add_run("_________________________________________________________________________________")
    r_div.font.color.rgb = RGBColor(220, 220, 220)

    # 2. Executive Abstract
    add_heading_styled(doc, "Executive Abstract", level=1)
    
    p1 = doc.add_paragraph(
        "Too many executive boards, risk officers, and internal auditors treat SOC 2 as a checkbox 'certification.' "
        "It is not. It is a highly specific, custom-scoped attestation report under AICPA standards. "
        "Often, organizations accept a vendor's SOC 2 report without realizing that the vendor has excluded their core "
        "service from the audit boundary, or that the report contains severe control failures ('exceptions') in Section IV, "
        "or that the validity of the report rests entirely on controls that the buying organization was supposed to implement "
        "but never did."
    )
    
    p2 = doc.add_paragraph(
        "This guide is designed for Internal Audit directors, risk professionals, and corporate leaders. "
        "It cuts through the technical jargon, exposes common misconceptions, provides a line-by-line reading guide to identify "
        "hidden risks, and delivers two actionable, enterprise-grade checklists:\n"
        "1. The 5-Phase Internal SOC 2 Readiness Checklist (to prepare your own company for an audit).\n"
        "2. The 10-Step Third-Party Vendor SOC 2 Review Checklist (to ingest and govern vendor risk)."
    )

    # 3. Part 1: What SOC 2 Actually Is
    add_heading_styled(doc, "Part 1: What SOC 2 Actually Is (and the Lies Vendors Tell You)", level=1)
    
    doc.add_paragraph(
        "To lead risk governance, you must first dismantle three prevailing myths:"
    )
    
    h3_1 = doc.add_paragraph()
    h3_1.paragraph_format.space_before = Pt(8)
    h3_1.paragraph_format.space_after = Pt(2)
    h3_1.paragraph_format.keep_with_next = True
    r_h3_1 = h3_1.add_run("Myth 1: 'Our vendor is SOC 2 Certified.'")
    r_h3_1.font.bold = True
    doc.add_paragraph(
        "The Reality: There is no such thing as a 'SOC 2 Certification.' An audit firm does not certify a company; "
        "they issue an attestation opinion regarding whether the company's description of their system is fairly presented, "
        "whether the controls are suitably designed, and (for Type II) whether those controls operated effectively over a specific period. "
        "A SOC 2 report is an opinion, and like all opinions, it can be clean, qualified, or worse."
    )

    h3_2 = doc.add_paragraph()
    h3_2.paragraph_format.space_before = Pt(8)
    h3_2.paragraph_format.space_after = Pt(2)
    h3_2.paragraph_format.keep_with_next = True
    r_h3_2 = h3_2.add_run("Myth 2: 'A SOC 2 Type I is just as good as a Type II.'")
    r_h3_2.font.bold = True
    doc.add_paragraph(
        "The Reality: A Type I report is a 'snapshot' of controls as of a single day (e.g., December 31). It proves that the "
        "controls exist on paper and are designed correctly, but it does not test whether they actually work in practice. "
        "A Type II report evaluates the operating effectiveness of controls over a testing window—typically 3 to 12 months. "
        "Accepting a Type I report for a critical, high-volume cloud provider is a massive governance gap; you are trusting a "
        "system based on a single day's performance."
    )

    h3_3 = doc.add_paragraph()
    h3_3.paragraph_format.space_before = Pt(8)
    h3_3.paragraph_format.space_after = Pt(2)
    h3_3.paragraph_format.keep_with_next = True
    r_h3_3 = h3_3.add_run("Myth 3: 'All SOC 2 reports cover the same things.'")
    r_h3_3.font.bold = True
    doc.add_paragraph(
        "The Reality: Every SOC 2 report is unique. While they are built upon the AICPA Trust Services Criteria (TSC), "
        "the service organization decides which criteria to include (beyond the mandatory Security criteria) and defines "
        "the boundaries of the system. A software vendor might have a SOC 2 report that only covers their corporate office HR "
        "processes, while excluding their actual SaaS hosting infrastructure."
    )

    # TSC Table
    add_heading_styled(doc, "The 5 Trust Services Criteria (TSC) Decoded", level=2)
    tsc_table = doc.add_table(rows=6, cols=2)
    tsc_table.style = 'Light Shading Accent 1'
    
    headers = ["CRITERIA", "DESCRIPTION & AUDIT FOCUS"]
    for i, name in enumerate(headers):
        cell = tsc_table.cell(0, i)
        cell.text = name
        set_cell_background(cell, "1D3557")
        set_cell_margins(cell, top=140, bottom=140, left=150, right=150)
        p = cell.paragraphs[0]
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(9.5)

    tsc_data = [
        ("Security\n(Mandatory Baseline)", "The 'Common Criteria'. Focuses on firewalls, logical access, multi-factor authentication (MFA), vulnerability scanning, intrusion detection, and physical data center protections. Protects against unauthorized access."),
        ("Availability", "Evaluates system uptime, performance monitoring, disaster recovery plans, data backup processes, and incident response capacity. Crucial for mission-critical SaaS, infrastructure, and hosting providers."),
        ("Processing Integrity", "Ensures system processing is complete, valid, accurate, timely, and authorized. Essential for financial technology (fintech), e-commerce platforms, payment gateways, and transactional clearinghouses."),
        ("Confidentiality", "Protects data designated as confidential (e.g., intellectual property, source code, corporate financials, pre-launch metrics) from unauthorized disclosure. Focuses on data classification and encryption in transit/at rest."),
        ("Privacy", "Governs the collection, use, retention, disclosure, and disposal of Personal Identifiable Information (PII) under regulatory compliance frameworks like GDPR, HIPAA, or the KSA Personal Data Protection Law (PDPL).")
    ]

    for row_idx, (crit, desc) in enumerate(tsc_data, start=1):
        cell_c = tsc_table.cell(row_idx, 0)
        cell_d = tsc_table.cell(row_idx, 1)
        cell_c.text = crit
        cell_d.text = desc
        
        # Apply style/borders/paddings
        for c in (cell_c, cell_d):
            set_cell_margins(c, top=120, bottom=120, left=150, right=150)
            p = c.paragraphs[0]
            p.runs[0].font.size = Pt(9.5)
        cell_c.paragraphs[0].runs[0].font.bold = True
        set_cell_background(cell_c, "F5F7FA")

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(12)

    # 4. Part 2: Anatomy of a SOC 2 Report
    add_heading_styled(doc, "Part 2: The SOC 2 Anatomy: A Line-by-Line Reading Guide", level=1)
    
    doc.add_paragraph(
        "A standard SOC 2 Type II report contains five distinct sections. As an Internal Auditor or Executive, "
        "you must know exactly where to look to find the high-risk entries. Do not read the report from front to back; "
        "jump directly to the sections outlined below."
    )

    add_heading_styled(doc, "1. Section I: Independent Service Auditor's Report (The Opinion)", level=2)
    doc.add_paragraph(
        "This is the auditor's legal attestation. Skip directly to the 'Opinion' paragraph. You will find one of four options:\n"
        "• Unmodified (Clean): The auditor agrees that the description is fair, the controls are designed correctly, and they operated effectively during the testing period.\n"
        "• Qualified: A major red flag. The auditor found that certain controls were either poorly designed or failed consistently, meaning one or more Trust Services Criteria were not fully met. Action: Read the qualification paragraph immediately to identify the failure (e.g., 'Except for the operating effectiveness of change management controls...').\n"
        "• Adverse: The control system is broken and cannot be relied upon. Do not onboard this vendor.\n"
        "• Disclaimer of Opinion: The auditor was unable to complete the audit due to a lack of evidence or cooperation. A critical warning."
    )

    add_heading_styled(doc, "2. Section III: Description of the System (The Boundary)", level=2)
    doc.add_paragraph(
        "This section is written by the vendor's management, describing their software, people, infrastructure, and procedures. You must verify:\n"
        "• Audit Scope: Does the description match the specific service you are buying? If you are purchasing their 'Enterprise API Analytics' platform, but the description only covers their 'Core SaaS Database Hosting,' you have a scope mismatch.\n"
        "• Subservice Organizations & CSOCs: Did the vendor outsource their hosting (e.g., to AWS or Microsoft Azure)? If so, did the auditor use the Carve-out Method or the Inclusive Method?\n"
        "  - Carve-out Method (Standard): The auditor excluded AWS controls from the audit, meaning you must obtain AWS's independent SOC 2 report separately to cover that risk layer.\n"
        "  - Inclusive Method: The auditor actually tested the controls at AWS as part of this audit (extremely rare).\n"
        "  - The CSOC Trap: When the vendor carves out AWS, AWS's SOC 2 report imposes CUECs on the vendor! These are called Complementary Subservice Organization Controls (CSOCs). As a buyer, you must look in Section III to ensure that the vendor has actively documented and mapped how they comply with AWS's CSOCs. If they ignored AWS's rules, their hosting layer remains highly vulnerable."
    )

    add_heading_styled(doc, "3. The CUEC Trap (Complementary User Entity Controls)", level=2)
    doc.add_paragraph(
        "Hidden at the end of Section III is the most dangerous element of a SOC 2 report: Complementary User Entity Controls (CUECs). "
        "Service organizations operate in a shared-responsibility model. The auditor writes the report assuming that YOU (the customer) are executing specific controls. "
        "If you fail to implement these CUECs, the vendor's SOC 2 report becomes effectively useless.\n"
        "• Example: A SaaS vendor's SOC 2 shows 100% encryption and secure multi-tenant isolation. However, under CUECs, they state: 'User entities are responsible for managing access credentials, enforcing multi-factor authentication (MFA) for administrative users, and performing quarterly access reviews.'\n"
        "• The Audit Gap: If your IT team has not set up single sign-on (SSO), has not enforced MFA, or has neglected quarterly reviews, a malicious actor can compromise an account easily. If a breach occurs, the vendor is legally protected because you failed to implement the CUECs listed in their SOC 2 report."
    )

    add_heading_styled(doc, "4. Section IV: Description of Tests of Controls and Results (The Exceptions)", level=2)
    doc.add_paragraph(
        "This is the technical core of the report, presenting a detailed table of every control activity tested by the auditor.\n"
        "• The 'Exceptions' Column: Scan this column for anything other than 'None.'\n"
        "• Evaluating Exceptions: If an exception is found (e.g., 'For a sample of 25 termination events, 3 users did not have their active access revoked within the required 24-hour window'), you must perform a quantitative risk assessment:\n"
        "  1. What was the sample size? (3 failures out of 25 is a 12% failure rate).\n"
        "  2. Did the auditor note any compensating controls?\n"
        "  3. How long did the terminated employees retain access?\n"
        "  4. Was this specific failure the reason for a qualified opinion in Section I?\n"
        "• Management's Response (Section V): If Section IV has exceptions, the vendor will explain them in Section V. Be highly skeptical. Look for actual root-cause corrections rather than generic excuses (e.g., 'It was a manual oversight during a holiday week')."
    )

    # 5. Part 3: Internal Readiness Checklist
    add_heading_styled(doc, "Part 3: The 5-Phase Internal SOC 2 Readiness Checklist", level=1)
    
    doc.add_paragraph(
        "For organizations building their own SOC 2 compliant environment from scratch. "
        "This checklist is structured to guide your internal audit team and control owners through an efficient, zero-waste preparation lifecycle, moving systematically from scoping to audit defense."
    )

    # Create Phase Table
    p_table = doc.add_table(rows=1, cols=5)
    p_table.style = 'Light Shading Accent 1'
    p_headers = ["REF", "PHASE / ACTION ITEM", "RESPONSIBILITY", "OWNER", "STATUS / EVIDENCE RECORD"]
    
    for i, name in enumerate(p_headers):
        cell = p_table.cell(0, i)
        cell.text = name
        set_cell_background(cell, "1D3557")
        set_cell_margins(cell, top=140, bottom=140, left=100, right=100)
        p = cell.paragraphs[0]
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(8.5)

    checklist_items = [
        # Phase 1
        ("Phase 1", "Scoping & Boundary Definition", "", "", ""),
        ("1.1", "Data Mapping: Map exactly where customer data (PII, IP, financials) enters, is processed, and exits your system.", "Security / CISO", "[ ] Not Started", "Data flow diagrams"),
        ("1.2", "System Boundary Definition: Document the technical infrastructure, software applications, people, and procedures in scope.", "Product / CTO", "[ ] Not Started", "System Boundary Document"),
        ("1.3", "TSC Selection: Select Security (Common Criteria) + Availability/Processing Integrity/Confidentiality/Privacy based on business SLAs.", "Internal Audit / CAE", "[ ] Not Started", "Signed Scoping Charter"),
        ("1.4", "Vendor Carve-Out Strategy: List all third-party subservice organizations (AWS, Stripe) and decide to carve them out of your report.", "IT / Engineering", "[ ] Not Started", "Vendor Inventory list"),
        
        # Phase 2
        ("Phase 2", "Control Design & Gap Analysis", "", "", ""),
        ("2.1", "Gap Assessment: Compare existing practices against selected Trust Services Criteria points. Document every gap.", "Internal Audit", "[ ] Not Started", "Gap Assessment Matrix"),
        ("2.2", "Policy Formalization: Draft and obtain board/management approval for InfoSec, Access Control, Change Management, and DR policies.", "Compliance / Risk", "[ ] Not Started", "Formal Policy PDFs"),
        ("2.3", "Control Activity Mapping: Design unique control activities per criterion stating who does what, how often, and the evidence generated.", "Internal Audit", "[ ] Not Started", "Controls Mapping Matrix"),
        
        # Phase 3
        ("Phase 3", "Implementation & Evidence Automation", "", "", ""),
        ("3.1", "Logical Access Controls: Enforce SSO/MFA on all critical/prod systems. Document legacy exceptions and enforce strict compensating controls (IP whitelisting, bastion hosts).", "IT / DevOps", "[ ] Not Started", "SSO/MFA config & exception logs"),
        ("3.2", "Evidence Logging Automation: Configure ticketing systems (Jira/ServiceNow) to link code commits directly to approved staging changes.", "DevOps / Eng", "[ ] Not Started", "GitHub / CI-CD pipeline config"),
        ("3.3", "Continuous Control Monitoring (CCM): Deploy a continuous GRC/monitoring platform or write automated monitoring alert scripts.", "Security Team", "[ ] Not Started", "GRC Dashboard / alert logs"),
        
        # Phase 4
        ("Phase 4", "The Dry Run (Pre-assessment testing)", "", "", ""),
        ("4.1", "Population Integrity Testing: Test the completeness and accuracy of evidence logs (active user rosters, code deploy histories).", "Internal Audit", "[ ] Not Started", "IA Population Test Sheet"),
        ("4.2", "Sample Testing & System Rules: Draw samples of 25 for manual controls. For automated rules (firewall configs), test the design rule directly (sample of 1).", "Internal Audit", "[ ] Not Started", "IA Audit Workpapers"),
        ("4.3", "Tabletop Simulations: Conduct formal, documented disaster recovery simulations and mock incident response drills.", "DevOps / Security", "[ ] Not Started", "DR Test & Incident Log report"),
        
        # Phase 5
        ("Phase 5", "Audit Window & Auditor Management", "", "", ""),
        ("5.1", "Audit Window Definition: Define the Type II testing period. (Note: While a 3-month window is used as a fast-track bridge, 6-month is the industry standard minimum).", "Internal Audit / CAE", "[ ] Not Started", "Engagement Letter"),
        ("5.2", "SPOC Establishment: Designate a Single Point of Contact (SPOC) between your firm and the CPA auditors to vet all evidence submissions.", "Internal Audit", "[ ] Not Started", "Communication Charter"),
        ("5.3", "Vetting & Audit Defense: Establish quality review of all submitted files. Confirm date ranges and verify zero unrelated client data leaks.", "Internal Audit / CAE", "[ ] Not Started", "Vetted Evidence Folder")
    ]

    for ref, action, resp, owner, evidence in checklist_items:
        row_cells = p_table.add_row().cells
        row_cells[0].text = ref
        row_cells[1].text = action
        row_cells[2].text = resp
        row_cells[3].text = owner
        row_cells[4].text = evidence
        
        is_phase_header = (ref.startswith("Phase") and action != "")
        
        # Style
        for idx, cell in enumerate(row_cells):
            set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
            p = cell.paragraphs[0]
            if len(p.runs) > 0:
                p.runs[0].font.size = Pt(8.5)
                if is_phase_header:
                    p.runs[0].font.bold = True
                    p.runs[0].font.color.rgb = RGBColor(163, 58, 33)
            
            if is_phase_header:
                set_cell_background(cell, "F5F7FA")
        
        # Special styling for Reference cell
        row_cells[0].paragraphs[0].runs[0].font.bold = True

    p_space2 = doc.add_paragraph()
    p_space2.paragraph_format.space_before = Pt(12)

    # 6. Part 4: Vendor SOC 2 Review Checklist
    add_heading_styled(doc, "Part 4: The 10-Step Third-Party Vendor SOC 2 Review Checklist", level=1)
    
    doc.add_paragraph(
        "For Risk and Internal Audit Professionals conducting third-party vendor risk assessments. "
        "When onboarding critical software vendors or SaaS providers, deploy this 10-step ingestion review scorecard. "
        "Never accept a SOC 2 logo at face value."
    )

    # Create Scorecard Table
    s_table = doc.add_table(rows=1, cols=4)
    s_table.style = 'Light Shading Accent 1'
    s_headers = ["STEP", "REVIEW REQUIREMENT", "COMPENSATING CONTROLS / RESIDUAL RISK", "OWNER / STATUS"]
    
    for i, name in enumerate(s_headers):
        cell = s_table.cell(0, i)
        cell.text = name
        set_cell_background(cell, "1D3557")
        set_cell_margins(cell, top=140, bottom=140, left=100, right=100)
        p = cell.paragraphs[0]
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(8.5)

    steps_data = [
        ("Step 1", "Report Freshness Verification: Ensure testing period ended within last 12 months. If older, obtain signed Bridge/Gap Letter.", "", "[ ] Pending Review"),
        ("Step 2", "Standard & CPA Firm Check: Verify issued by licensed AICPA CPA firm in good standing (AICPA Peer Review Directory). Watch out for lookalikes.", "", "[ ] Pending Review"),
        ("Step 3", "Type II vs. Type I: Confirm Type II (operating effectiveness window). If Type I, log as high risk and request Type II roadmap.", "", "[ ] Pending Review"),
        ("Step 4", "Scope & Boundary Alignment: Read Section III system description. Confirm that the specific software/API/data region purchased is in scope.", "", "[ ] Pending Review"),
        ("Step 5", "TSC Sufficiency: Verify correct criteria are audited (e.g. Availability for cloud kitchen logistics; Privacy for personal PII).", "", "[ ] Pending Review"),
        ("Step 6", "Auditor's Opinion Check: Confirm Section I is Unmodified (Clean). If Qualified, map failures to contract risk profile.", "", "[ ] Pending Review"),
        ("Step 7", "Subservice & CSOCs Assessment: Identify carved-out providers (AWS). Verify vendor maps and meets AWS's CSOC (Subservice CUEC) requirements.", "", "[ ] Pending Review"),
        ("Step 8", "CUEC Extraction & Mapping: Extract all Complementary User Entity Controls from Section III. Assign internal owners to implement.", "", "[ ] Pending Review"),
        ("Step 9", "Quantitative Exception Analysis: Go to Section IV. Calculate failure rates of exceptions. Assess remediation speed and impact.", "", "[ ] Pending Review"),
        ("Step 10", "Contract Alignment & Residual Risk: Log unmitigated gaps in Risk Register. Add security SLAs and liability caps in MSA if critical.", "", "[ ] Pending Review")
    ]

    for step, req, comp, owner in steps_data:
        row_cells = s_table.add_row().cells
        row_cells[0].text = step
        row_cells[1].text = req
        row_cells[2].text = comp
        row_cells[3].text = owner
        
        for cell in row_cells:
            set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
            p = cell.paragraphs[0]
            if len(p.runs) > 0:
                p.runs[0].font.size = Pt(8.5)
        
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        set_cell_background(row_cells[0], "F5F7FA")

    p_space3 = doc.add_paragraph()
    p_space3.paragraph_format.space_before = Pt(12)

    # 7. Conclusion
    add_heading_styled(doc, "Conclusion", level=1)
    doc.add_paragraph(
        "SOC 2 is not a certificate of invincibility. It is a structured transparency report. "
        "As an internal audit professional, your objective is to read between the lines: find the boundaries, identify the exceptions in Section IV, "
        "map your company's responsibilities in the CUECs, and establish a continuous, data-driven controls environment. "
        "By shifting from historical sampling to 100% population testing and automating evidence capture, you turn SOC 2 from an "
        "annual audit headache into a robust, real-time risk shield for your enterprise."
    )

    # Save Document
    target_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "public", "downloads"))
    os.makedirs(target_dir, exist_ok=True)
    target_path = os.path.join(target_dir, "SOC_2_Internal_Audit_Readiness_and_Review_Guide.docx")
    
    print(f"Saving document to {target_path}...")
    doc.save(target_path)
    print("✓ Successfully generated DOCX file!")

if __name__ == '__main__':
    main()
